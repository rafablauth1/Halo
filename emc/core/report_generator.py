from datetime import datetime, timezone

from docx import Document

from emc.config import REPORTS_DIR, STANDARDS
from emc.core.db import db_cursor
from emc.core.planner import get_project
from emc.core.test_session import get_session, list_events


def generate_report(session_id: int) -> str:
    session = get_session(session_id)
    if session is None:
        raise ValueError(f"Sessão {session_id} não encontrada")
    project = get_project(session["project_id"])
    events = list_events(session_id)

    doc = Document()
    doc.add_heading("Laudo de Ensaio EMC", level=1)

    doc.add_heading("Identificação", level=2)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light List"
    rows = [
        ("Projeto", project["name"] if project else "-"),
        ("Cliente", project["client"] if project else "-"),
        ("Norma", f"{session['standard_code']} — {STANDARDS.get(session['standard_code'], '')}"),
        ("Nível/Parâmetro", session["level_label"] or "-"),
        ("EUT", session["eut_name"] or "-"),
        ("Número de série do EUT", session["eut_serial"] or "-"),
        ("Operador", session["operator"] or "-"),
        ("Início", session["started_at"] or "-"),
        ("Término", session["finished_at"] or "-"),
        ("Resultado", (session["result"] or "pendente").upper()),
    ]
    for label, value in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = str(value)

    if session["notes"]:
        doc.add_heading("Observações", level=2)
        doc.add_paragraph(session["notes"])

    doc.add_heading("Registro do ensaio", level=2)
    if events:
        log_table = doc.add_table(rows=1, cols=2)
        log_table.style = "Light List"
        header = log_table.rows[0].cells
        header[0].text = "Horário"
        header[1].text = "Evento"
        for event in events:
            cells = log_table.add_row().cells
            cells[0].text = event["timestamp"]
            cells[1].text = event["message"]
    else:
        doc.add_paragraph("Nenhum evento registrado.")

    REPORTS_DIR.mkdir(exist_ok=True)
    filename = f"laudo_sessao_{session_id}_{session['standard_code'].replace('-', '')}.docx"
    file_path = REPORTS_DIR / filename
    doc.save(str(file_path))

    with db_cursor() as cur:
        cur.execute(
            "INSERT INTO reports (session_id, file_path, generated_at) VALUES (?, ?, ?)",
            (session_id, str(file_path), datetime.now(timezone.utc).isoformat()),
        )

    return str(file_path)


def list_reports_for_session(session_id: int) -> list[dict]:
    with db_cursor() as cur:
        cur.execute(
            "SELECT * FROM reports WHERE session_id = ? ORDER BY generated_at DESC",
            (session_id,),
        )
        return [dict(row) for row in cur.fetchall()]
